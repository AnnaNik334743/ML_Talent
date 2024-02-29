import datetime
import io
import os
import shutil
from typing import Optional

import docx
import docx2txt
import fitz
import numpy as np
from PIL import Image
from easyocr import Reader
from langdetect import detect
from spire.doc import FileFormat, Document
from json_schemas import TruncParserOutput, ParserOutput
from utils import naive_lang_detect, is_human
from main import EXTRACTED_IMG_FOLDER, BUCKET_NAME, S3


def save_if_img_contains_human(img: Image.Image) -> Optional[tuple[str, str]]:
    """
    Saves an image if it contains a human.

    Parameters:
        img (PIL.Image.Image): The image to be analyzed.
        local_img_path (str): The local folder to save extracted images.
        yandex_img_path (str): The remote folder to save extracted images (yandexcloud).

    Returns:
        bool: found_img_path if a human is detected in the image, None otherwise.
    """
    if is_human(img):
        local_img_path = EXTRACTED_IMG_FOLDER + '/' + datetime.datetime.now().strftime(
            "%Y-%m-%d_%H-%M-%S") + '_image.jpg'
        img.save(local_img_path)

        bucket_filename = local_img_path
        S3.upload_file(local_img_path, BUCKET_NAME, bucket_filename)
        yandex_img_path = f"https://storage.yandexcloud.net/{BUCKET_NAME}/{bucket_filename}"

        return local_img_path, yandex_img_path
    return None


def pdf_parser(file_path: str) -> 'TruncParserOutput':
    """
    Parses a PDF file.

    Parameters:
        file_path (str): The path of the PDF file to parse.

    Returns:
        TruncParserOutput: The parsed output.
    """
    yandex_img_path = ""

    doc = fitz.open(file_path)
    text = []

    for page_number, page in enumerate(doc):
        page = doc.load_page(page_number)

        # Extract text from each page
        text.append(page.get_text())

        # Extract images from each page
        image_list = page.get_images(full=True)

        for image_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_data = io.BytesIO(image_bytes)
            img = Image.open(image_data)

            paths = save_if_img_contains_human(img)
            if paths:
                local_img_path, yandex_img_path = paths
                break

    # If we have parsed too few words, we probably need to use Optical Character Recognition for parsing the file
    if len(''.join(text).strip().split()) < 80:

        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=300)
            img = np.ndarray([pix.h, pix.w, 3], dtype=np.uint8, buffer=pix.samples_mv)

            # On the first iteration (first page only), decide whether the text is in English or in Russian
            # for better parsing
            if i == 0:
                reader = Reader(['en', 'ru'])
                language = naive_lang_detect(' '.join([result[1] for result in reader.readtext(img)]))
                if language == 'en':
                    reader = Reader(["en"])

            # Parse the whole file using OCR
            text.append(' '.join([result[1] for result in reader.readtext(img)]))

    return TruncParserOutput(**{'photo_path': yandex_img_path, 'text': '\n'.join(text)})


def docx_parser(file_path: str) -> 'TruncParserOutput':
    """
    Parses a DOCX file.

    Parameters:
        file_path (str): The path of the DOCX file to parse.

    Returns:
        TruncParserOutput: The parsed output.
    """
    # When parsing DOCX, we need to use temp folder for storing images
    temp_img_folder = 'temp_img_folder'

    try:
        os.mkdir(temp_img_folder)
    except FileExistsError:
        pass

    yandex_img_path = ""

    # Extract images from file
    docx2txt.process(file_path, temp_img_folder)
    for imgfile in os.listdir(temp_img_folder):
        imgpath = temp_img_folder + '/' + imgfile
        img = Image.open(imgpath)

        paths = save_if_img_contains_human(img)
        if paths:
            local_img_path, yandex_img_path = paths
            break

    # Extract text from file
    doc = docx.Document(file_path)
    text = [paragraph.text for paragraph in doc.paragraphs]

    # Remove temp folder for images
    shutil.rmtree(temp_img_folder)

    # If we have parsed too few words, there is probably a table in the file, so we need to parse it as a table
    if len(''.join(text).strip().split()) < 80:
        text = ['\n'.join(['\n'.join([cell.text for cell in row.cells]) for row in table.rows]) for table in
                doc.tables]

    return TruncParserOutput(**{'photo_path': yandex_img_path, 'text': '\n'.join(text)})


def parse(file_path: str, file_extension: str) -> 'ParserOutput':
    """
    Parses a file based on its extension.

    Parameters:
        file_path (str): The path of the file to parse.
        file_extension (str): The extension of the file.

    Returns:
        ParserOutput: The parsed output.
    """

    if file_extension == 'pdf':
        data = pdf_parser(file_path)

    elif file_extension in ['doc', 'docx']:

        # Convert doc to docx for further processing
        if file_extension == 'doc':
            document = Document()
            document.LoadFromFile(file_path)
            file_path += 'x'
            document.SaveToFile(file_path, FileFormat.Docx2016)
            document.Close()

        data = docx_parser(file_path)
    else:
        raise ValueError("Unknown file extension!")

    try:
        # Get the language of the text
        language = detect(data.get('text'))
    except Exception:
        language = ''

    # Update parsed data with newly extracted fields
    data = data.__dict__
    data.update({'filepath': file_path, 'file_extension': file_extension, 'language': language})

    return ParserOutput(**data)
